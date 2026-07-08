"""Round-trip test through a real generated .docx.

The DOCX extractor carries two subtle, regression-prone fixes that plain .txt
fixtures cannot exercise:

* v9 FIX A — merged-cell criteria tables (cols 0&1 identical, 2&3 identical).
* v9 FIX B — cover tables with no colon separator between key and value.

We build such a document with python-docx and assert both survive extraction +
parsing. If python-docx is unavailable the test skips (never fails the suite).
"""

from __future__ import annotations

from pathlib import Path

import pytest

pytest.importorskip("docx")

from vpat_reviewer.extraction.docx import DocxExtractor  # noqa: E402
from vpat_reviewer.parsing.document import parse_document  # noqa: E402


def _build_merged_cell_docx(path: str) -> None:
    from docx import Document

    d = Document()

    # Cover table: 2-col key/value with NO colon (v9 FIX B).
    cover = d.add_table(rows=2, cols=2)
    cover.rows[0].cells[0].text = "Name of Product/Version"
    cover.rows[0].cells[1].text = "Merged Product 9.0"
    cover.rows[1].cells[0].text = "Report Date"
    cover.rows[1].cells[1].text = "March 2024"

    # Criteria table: 5-col merged layout (cols 0&1 identical, 2&3 identical) (v9 FIX A).
    t = d.add_table(rows=2, cols=5)
    row0 = t.rows[0].cells
    row0[0].text = row0[1].text = "1.4.3 Contrast (Minimum) (Level AA)"
    row0[2].text = row0[3].text = "Does Not Support"
    row0[4].text = "Low contrast on buttons."
    row1 = t.rows[1].cells
    row1[0].text = row1[1].text = "2.1.1 Keyboard (Level A)"
    row1[2].text = row1[3].text = "Supports"
    row1[4].text = ""

    d.save(path)


def test_docx_merged_cells_and_cover(tmp_path: Path) -> None:
    p = tmp_path / "merged.docx"
    _build_merged_cell_docx(str(p))

    raw = DocxExtractor().extract(str(p))
    # v9 FIX B: cover key/value injected as "Key: Value" so metadata is findable.
    assert "Name of Product/Version: Merged Product 9.0" in raw.text

    doc = parse_document(raw)
    assert doc.product_name == "Merged Product"
    assert doc.product_version == "9.0"

    wcag = {c.criterion_id: c for c in doc.criteria if c.section == "wcag"}
    # v9 FIX A: status read from the correct (merged) column, not "Not Evaluated".
    assert wcag["1.4.3"].normalized_status == "Does Not Support"
    assert wcag["1.4.3"].remarks == "Low contrast on buttons."
    assert wcag["2.1.1"].normalized_status == "Supports"
