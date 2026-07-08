"""DOCX extractor (python-docx). Behavior relocated verbatim from the v10 parser.

Carries two hard-won fixes for merged-cell VPAT layouts (v9 FIX A) and colon-less
cover tables (v9 FIX B). Do not "simplify" the cell handling without a fixture.
"""

from __future__ import annotations

import logging

from vpat_reviewer.extraction.base import RawDocument, Table

logger = logging.getLogger(__name__)


class DocxExtractor:
    extensions: tuple[str, ...] = (".docx",)

    def extract(self, path: str) -> RawDocument:
        try:
            from docx import Document

            doc = Document(path)
            parts = [p.text for p in doc.paragraphs]
            tables: list[Table] = []
            for tbl in doc.tables:
                rows: Table = []
                for row in tbl.rows:
                    cells_raw = [cell.text.strip() for cell in row.cells]
                    n = len(cells_raw)

                    # Detect merged-cell VPAT layout:
                    # 5-col: [crit, crit_dup, status, status_dup, remarks]
                    #        -> normalise to [crit, status, remarks]
                    # 6-col: [crit, crit_dup, status, status_dup, remarks_dup, remarks]
                    #        -> same normalisation
                    # This avoids content-based deduplication which shifts column indices.
                    # v9 FIX A: DOCX merged-cell layout (e.g. Minitab VPAT). 5-col
                    # tables have cols 0&1 identical (criterion) and cols 2&3 identical
                    # (status). Deduplicating by content shifts indices, making status
                    # land in the wrong slot -> all "Not Evaluated". Fix: detect the
                    # merge and read cols 0, 2, 4 directly.
                    if n >= 5 and cells_raw[0] and cells_raw[0] == cells_raw[1]:
                        crit = cells_raw[0]
                        status = cells_raw[2] if n > 2 else ""
                        remarks = cells_raw[4] if n > 4 else (cells_raw[3] if n > 3 else "")
                        rows.append([crit, status, remarks])
                    else:
                        # Standard layout — keep as-is (no deduplication).
                        rows.append(list(cells_raw))

                tables.append(rows)
                # Also add a text representation for metadata extraction.
                parts.append(" | ".join(c.text.strip() for row in tbl.rows for c in row.cells[:2]))

            # v9 FIX B: DOCX cover tables use plain rows with no colon separator
            # (e.g. ["Name of Product/Version", "Minitab Online"]). Inject as
            # "Key: Value" so metadata regexes can find product name, date, etc.
            if doc.tables:
                cover_lines = []
                for row in doc.tables[0].rows:
                    cells = row.cells
                    if len(cells) >= 2:
                        key = cells[0].text.strip()
                        val = cells[1].text.strip()
                        if key and val:
                            cover_lines.append(f"{key}: {val}")
                if cover_lines:
                    parts.insert(0, "\n".join(cover_lines))

            return RawDocument("\n".join(parts), tables)
        except Exception as e:  # noqa: BLE001 — extraction failure must not crash.
            logger.warning("DOCX extraction error: %s", e)
            return RawDocument("", [])
