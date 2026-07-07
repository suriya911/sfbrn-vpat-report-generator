"""
VPAT Parser v5 — All audit findings corrected.
Handles PDF and DOCX. Fixes: standards detection, description contamination,
score formula (NA excluded from denominator), 602.3 status, remarks cleanup.
"""

import re
import logging
from datetime import date
from dataclasses import dataclass, field
from typing import Optional
from pathlib import Path

logger = logging.getLogger(__name__)


# ── Data models ───────────────────────────────────────────────────────────────

# Canonical definitions now live in the vpat_reviewer package (Phase 1 strangler);
# re-exported so the rest of this legacy module keeps working unchanged.
from vpat_reviewer.domain.models import VPATCriterion, VPATData


# ── Status normalisation ──────────────────────────────────────────────────────

from vpat_reviewer.domain.normalization import normalize_status


# ── Date parsing ──────────────────────────────────────────────────────────────

_MONTHS = {
    "january":1,"february":2,"march":3,"april":4,"may":5,"june":6,
    "july":7,"august":8,"september":9,"october":10,"november":11,"december":12,
    "jan":1,"feb":2,"mar":3,"apr":4,"jun":6,"jul":7,"aug":8,
    "sep":9,"sept":9,"oct":10,"nov":11,"dec":12,
}

def _parse_date(s: str) -> Optional[date]:
    if not s:
        return None
    s = re.sub(r'(\d+)(st|nd|rd|th)\b', r'\1', s, flags=re.IGNORECASE)
    patterns = [
        (r"(\w+)\s+(\d{1,2}),?\s+(\d{4})", "mname_d_y"),
        (r"(\d{1,2})\s+(\w+)\s+(\d{4})",   "d_mname_y"),
        (r"(\d{4})[/-](\d{1,2})[/-](\d{1,2})", "ymd"),
        (r"(\d{1,2})[/-](\d{1,2})[/-](\d{2,4})", "mdy"),
    ]
    for pat, fmt in patterns:
        m = re.search(pat, s, re.IGNORECASE)
        if not m:
            continue
        try:
            if fmt == "mname_d_y":
                mo = _MONTHS.get(m.group(1).lower())
                if mo:
                    return date(int(m.group(3)), mo, int(m.group(2)))
            elif fmt == "d_mname_y":
                mo = _MONTHS.get(m.group(2).lower())
                if mo:
                    return date(int(m.group(3)), mo, int(m.group(1)))
            elif fmt == "ymd":
                return date(int(m.group(1)), int(m.group(2)), int(m.group(3)))
            elif fmt == "mdy":
                yr = int(m.group(3))
                if yr < 100:
                    yr += 2000
                return date(yr, int(m.group(1)), int(m.group(2)))
        except (ValueError, TypeError):
            continue
    # v9 FIX D: month-year only dates (e.g. "August 2020") have no day number.
    # Convention: treat as 1st of the month. This allows outdated-VPAT detection
    # to work correctly for VPATs like Minitab (Aug 2020 = ~70 months old).
    # Month-year only (e.g. "August 2020") — treat as 1st of month
    # This is common in VPATs that only specify month and year
    my = re.search(r'(\w+)\s+(\d{4})\s*$', s.strip(), re.IGNORECASE)
    if my:
        mo = _MONTHS.get(my.group(1).lower())
        if mo:
            try:
                return date(int(my.group(2)), mo, 1)
            except (ValueError, TypeError):
                pass
    return None


# ── Text extraction ───────────────────────────────────────────────────────────

def _extract_pdf(path: str):
    try:
        import pdfplumber
        pages_text, all_tables = [], []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    pages_text.append(t)
                tbls = page.extract_tables()
                if tbls:
                    all_tables.extend(tbls)
        return "\n".join(pages_text), all_tables
    except Exception as e:
        logger.warning(f"PDF extraction error: {e}")
        return "", []


def _extract_docx(path: str):
    try:
        from docx import Document
        doc = Document(path)
        parts = [p.text for p in doc.paragraphs]
        tables = []
        for tbl in doc.tables:
            rows = []
            for row in tbl.rows:
                cells_raw = [cell.text.strip() for cell in row.cells]
                n = len(cells_raw)

                # Detect merged-cell VPAT layout:
                # 5-col: [crit, crit_dup, status, status_dup, remarks]
                #        → normalise to [crit, status, remarks]
                # 6-col: [crit, crit_dup, status, status_dup, remarks_dup, remarks]
                #        → same normalisation
                # This avoids content-based deduplication which shifts column indices.
                # v9 FIX A: DOCX merged-cell layout (e.g. Minitab VPAT)
                # 5-col tables have cols 0&1 identical (criterion) and cols 2&3
                # identical (status). Deduplicating by content shifts indices,
                # making status land in the wrong slot → all "Not Evaluated".
                # Fix: detect the merge and read cols 0, 2, 4 directly.
                if n >= 5 and cells_raw[0] and cells_raw[0] == cells_raw[1]:
                    # Merged criterion column detected
                    crit    = cells_raw[0]
                    status  = cells_raw[2] if n > 2 else ""
                    remarks = cells_raw[4] if n > 4 else (cells_raw[3] if n > 3 else "")
                    rows.append([crit, status, remarks])
                else:
                    # Standard layout — keep as-is (no deduplication)
                    rows.append(cells_raw)

            tables.append(rows)
            # Also add text representation for metadata extraction
            parts.append(" | ".join(c.text.strip() for row in tbl.rows for c in row.cells[:2]))

        # v9 FIX B: DOCX cover tables use plain rows with no colon separator
        # (e.g. ["Name of Product/Version", "Minitab Online"]) rather than
        # "Name of Product/Version: Minitab Online". Inject as "Key: Value"
        # so _extract_meta regexes can find product name, date, contact, etc.
        # Extra pass: extract cover table key/value pairs as "Key: Value\n" text
        # so _extract_meta can find product name, date, etc.
        if doc.tables:
            cover_lines = []
            for row in doc.tables[0].rows:
                cells = row.cells
                if len(cells) >= 2:
                    key = cells[0].text.strip()
                    val = cells[1].text.strip()
                    if key and val:
                        cover_lines.append(f"{key}: {val}")
            if cover_lines:
                parts.insert(0, "\n".join(cover_lines))

        return "\n".join(parts), tables
    except Exception as e:
        logger.warning(f"DOCX extraction error: {e}")
        return "", []


def _extract_txt(path: str):
    try:
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            return f.read(), []
    except Exception as e:
        logger.warning(f"TXT read error: {e}")
        return "", []


# ── Text cleanup helpers ──────────────────────────────────────────────────────

# Watermark/footer text that leaks into extracted text
_WATERMARK_PATTERNS = [
    re.compile(r'"Voluntary Product Accessibility Template"[^\n]*\n[^\n]*Page\s+\d+\s+of\s+\d+', re.IGNORECASE),
    re.compile(r'Page\s+\d+\s+of\s+\d+', re.IGNORECASE),
    re.compile(r'_{10,}'),   # long underscores (horizontal rules)
]

def _clean_extracted_text(text: str) -> str:
    """Remove watermarks, page numbers, and boilerplate from extracted text."""
    for pat in _WATERMARK_PATTERNS:
        text = pat.sub(' ', text)
    return text

def _is_blank_or_garbage(text: str) -> bool:
    """
    Return True if a field value is blank, a VPAT section header,
    or garbage text that should be treated as missing.
    """
    if not text or not text.strip():
        return True
    t = text.strip().lower()
    # Starts with known boilerplate
    garbage_starts = [
        "contact information", "notes:", "note:", "see wcag",
        "heading cell", "voluntary product", '"voluntary',
        "this report", "the testing", "remark", "n/a",
    ]
    for g in garbage_starts:
        if t.startswith(g):
            return True
    # Very short but non-meaningful
    if len(t) < 3:
        return True
    return False


# ── Metadata extraction ───────────────────────────────────────────────────────

# Known company names — product-specific only, excludes browser/tool names
_KNOWN_VENDORS = [
    "Atlassian", "Microsoft", "Instructure", "Blackboard", "Anthology",
    "Adobe", "Salesforce", "Respondus", "D2L", "Brightspace", "Moodle",
    "Kaltura", "Panopto", "Echo360", "Turnitin", "ProctorU", "Honorlock",
    "McGraw-Hill", "Pearson", "Cengage", "VitalSource", "Chegg",
    "LinkedIn", "Slack", "Trello",
]

def _extract_meta(text: str) -> dict:
    meta = {}

    # Product name / version
    for pat in [
        r"Name\s+of\s+Product[/\\]?Version\s*:\s*(.+)",
        r"Product\s+Name\s*:\s*(.+)",
        r"Product\s*[:/]\s*(.+)",
    ]:
        m = re.search(pat, text, re.IGNORECASE)
        if m:
            raw = m.group(1).strip().split("\n")[0].strip()
            # Try version with v/Version prefix first (e.g. "App v2.0" or "App Version 2.0")
            vm = re.search(r"[\s,;]+[Vv](?:ersion)?\s*([\d.]+)", raw)
            if vm:
                meta["product_name"] = _clean_product_name(raw[:vm.start()].strip())
                meta["product_version"] = vm.group(1)
            else:
                # Try bare decimal version at end of string (e.g. "TestApp 1.0")
                # Requires at least one decimal point to avoid splitting "Windows 10" → ["Windows","10"]
                vm2 = re.search(r"[\s]+([\d]+\.[\d]+(?:\.[\d]+)*)\s*$", raw)
                if vm2:
                    meta["product_name"] = _clean_product_name(raw[:vm2.start()].strip())
                    meta["product_version"] = vm2.group(1)
                else:
                    meta["product_name"] = _clean_product_name(raw)
            break

    # Report date — handle ordinal suffixes and simple "Date:" label
    for pat in [
        r"Report\s+Date\s*:\s*(.+)",
        r"Date\s+of\s+Report\s*:\s*(.+)",
        r"Evaluation\s+Date\s*:\s*(.+)",
        r"^Date\s*:\s*(.+)",                # v9 FIX C: bare "Date:" label (no "Report")
        r"(?:(?!Report\s|Evaluation\s)\b)Date\s*:\s*(.+)",  # "Date:" not preceded by another word
    ]:
        m = re.search(pat, text, re.IGNORECASE)
        if m:
            raw_date = m.group(1).strip().split("\n")[0].strip()
            if not _is_blank_or_garbage(raw_date):
                meta["vendor_report_date_raw"] = raw_date
                break

    # Vendor — explicit field
    for pat in [
        r"(?:Company|Vendor|Manufacturer)\s*(?:Name)?\s*:\s*(.+)",
    ]:
        m = re.search(pat, text, re.IGNORECASE)
        if m:
            v = m.group(1).strip().split("\n")[0].strip()
            if v and len(v) < 80 and not _is_blank_or_garbage(v):
                meta["vendor_name"] = v
                break

    # Vendor — scan HEADER region (first 20%) for known company names only
    if not meta.get("vendor_name"):
        header = text[:max(int(len(text) * 0.20), 800)]
        for vendor in _KNOWN_VENDORS:
            if re.search(r'\b' + re.escape(vendor) + r'\b', header, re.IGNORECASE):
                meta["vendor_name"] = vendor
                break

    # Vendor from email domain (last resort)
    if not meta.get("vendor_name"):
        em = re.search(r'[\w._%+-]+@([\w]+)\.', text)
        if em:
            domain = em.group(1).capitalize()
            if domain.lower() not in ("gmail", "yahoo", "hotmail", "outlook",
                                      "google", "microsoft", "apple"):
                meta["vendor_name"] = domain

    # VPAT edition
    m = re.search(r"VPAT[®\s]*Version\s*([\d.]+)", text, re.IGNORECASE)
    if m:
        meta["vpat_edition"] = f"VPAT\u00ae Version {m.group(1)}"

    # Contact info — must look like actual contact data (email, phone, name)
    m = re.search(r"Contact\s+Information\s*:\s*(.+?)(?:\n\n|\Z)", text, re.IGNORECASE | re.DOTALL)
    if m:
        c = m.group(1).strip().split("\n")[0].strip()
        # Only keep if looks like actual contact info
        if (c and len(c) < 150
                and not _is_blank_or_garbage(c)
                and re.search(r'[@.\w]{4,}', c)):
            meta["vendor_contact"] = c

    # Email as contact fallback
    if not meta.get("vendor_contact"):
        em = re.search(r'[\w._%+-]+@[\w.-]+\.[A-Z]{2,}', text, re.IGNORECASE)
        if em:
            meta["vendor_contact"] = em.group(0)

    # v9 FIX E: description regex must stop at the next cover-table key name.
    # Without this, the regex captured "Introductory Minitab statistics package...\n
    # Date: August  2020" because it ran past the description into the date row.
    # Product description — stop at any next cover-table key line
    m = re.search(
        r"Product\s+Description\s*:\s*(.+?)(?:\n(?:Date|Contact|Notes|Evaluation|Name|Version|Vendor)\s*:|\nContact|\nNotes|\nEvaluation|\n\n|\Z)",
        text, re.IGNORECASE | re.DOTALL)
    if m:
        d = m.group(1).strip()
        if not _is_blank_or_garbage(d) and len(d) > 30:
            d = _clean_extracted_text(d)[:800]
            meta["product_description"] = d

    # Evaluation methods — strip watermarks
    m = re.search(
        r"Evaluation\s+Methods?\s+Used\s*:\s*(.+?)(?:\n(?:Applicable|Standards|Table\s+\d|Terms)\b|\Z)",
        text, re.IGNORECASE | re.DOTALL
    )
    if m:
        methods = _clean_extracted_text(m.group(1).strip())
        # Remove ALL watermark and ITI boilerplate lines
        # Pattern covers: "Voluntary Product Accessibility Template" and "VPAT" are registered
        methods = re.sub(r'"Voluntary Product Accessibility Template"[^\n]*\n?', '', methods)
        methods = re.sub(r'[^\n]*VPAT[^\n]*registered[^\n]*\n?', '', methods, flags=re.IGNORECASE)
        methods = re.sub(r'service marks of[^\n]*\n?', '', methods, flags=re.IGNORECASE)
        methods = re.sub(r'Information Technology Industry[^\n]*\n?', '', methods, flags=re.IGNORECASE)
        methods = re.sub(r'Page\s+\d+\s+of\s+\d+[^\n]*\n?', '', methods, flags=re.IGNORECASE)
        methods = re.sub(r'\n{3,}', '\n\n', methods).strip()
        if methods and len(methods) > 20:
            meta["evaluation_methods"] = methods[:2000]

    # Product type
    if re.search(r"web\s*app|web\s*portal|website|software", text, re.IGNORECASE):
        meta["product_type"] = "Web Application / Software"

    return meta


def _clean_product_name(raw: str) -> str:
    if not raw:
        return raw
    cleaned = re.sub(r'\s+\d{4}\s+Q\d\b', '', raw, flags=re.IGNORECASE).strip()
    cleaned = re.sub(r'\s+VPAT\b.*$', '', cleaned, flags=re.IGNORECASE).strip()
    cleaned = re.sub(r'\s+v[\d.]+\s*$', '', cleaned, flags=re.IGNORECASE).strip()
    cleaned = re.sub(r'\s+\d{4}-\d{2}-\d{2}.*$', '', cleaned).strip()
    return cleaned if cleaned else raw


# ── Standards detection — FIX: support both 'WCAG' and 'Web Content Accessibility Guidelines' ──

def _detect_standards(text: str) -> list:
    """
    Detect standards from VPATs that spell out 'Web Content Accessibility Guidelines'
    (not just the abbreviation 'WCAG').
    """
    prefix = r"(?:WCAG|Web\s+Content\s+Accessibility\s+Guidelines)\s*"
    checks = [
        (prefix + r"2\.0.*?Level\s*A(?!\s*A)",     "WCAG 2.0 Level A"),
        (prefix + r"2\.0.*?Level\s*AA\b",           "WCAG 2.0 Level AA"),
        (prefix + r"2\.1.*?Level\s*A(?!\s*A)",      "WCAG 2.1 Level A"),
        (prefix + r"2\.1.*?Level\s*AA\b",           "WCAG 2.1 Level AA"),
        (prefix + r"2\.2.*?Level\s*A(?!\s*A)",      "WCAG 2.2 Level A"),
        (prefix + r"2\.2.*?Level\s*AA\b",           "WCAG 2.2 Level AA"),
        (r"(?:Section\s*508|Revised\s*508)",         "Section 508 (Revised 2017)"),
        (r"EN\s*301\s*549",                          "EN 301 549"),
    ]
    standards = []
    seen = set()
    for pat, label in checks:
        if label not in seen and re.search(pat, text, re.IGNORECASE | re.DOTALL):
            standards.append(label)
            seen.add(label)
    return standards


# ── Criterion parsing helpers ─────────────────────────────────────────────────

_CRIT_RE = re.compile(
    r"(\d+\.\d+\.\d+)\s+(.+?)\s*\(Level\s+(A{1,3})\)",
    re.IGNORECASE
)
_STATUS_RE = re.compile(
    r"^(Supports\s+with\s+Exceptions|Partially\s+Supports?|Does\s+Not\s+Supports?|"
    r"Not\s+Applicable|Not\s+Evaluated|Supports?|Supported|Support\b)",
    re.IGNORECASE
)

def _clean_criterion_title(title: str) -> str:
    """Remove cross-reference boilerplate from criterion titles."""
    title = re.sub(r'\s+Also\s+applies\s+to\s*:.*$', '', title,
                   flags=re.IGNORECASE | re.DOTALL)
    title = re.sub(r'\s+Revised\s+Section\s+508.*$', '', title,
                   flags=re.IGNORECASE | re.DOTALL)
    title = re.sub(r'\s+EN\s+301\s+549.*$', '', title,
                   flags=re.IGNORECASE | re.DOTALL)
    title = re.sub(r'\s+\d{3}\s+\(Web\).*$', '', title,
                   flags=re.IGNORECASE | re.DOTALL)
    return title.strip()

def _clean_remarks(remarks: str) -> str:
    if not remarks:
        return ""
    r = re.sub(r"^Remark\s*s?\s*:", "", remarks, flags=re.IGNORECASE).strip()
    # Strip watermark text that bleeds in
    r = _clean_extracted_text(r)
    r = re.sub(r'"Voluntary Product[^"]*"[^\n]*\n?', '', r)
    r = re.sub(r'service marks of[^\n]*\n?', '', r)
    return r.strip()


# ── Table criterion parsing ───────────────────────────────────────────────────

def _parse_from_tables(tables: list) -> list:
    criteria = []
    seen = set()

    for table in tables:
        if not table or len(table) < 2:
            continue

        # Detect column layout
        col_layout = None
        for row in table:
            if not row:
                continue
            cell0 = str(row[0] or "").strip()
            if _CRIT_RE.search(cell0):
                if len(row) >= 9:
                    c3 = str(row[3] or "").strip() if len(row) > 3 else ""
                    c6 = str(row[6] or "").strip() if len(row) > 6 else ""
                    col_layout = "9col" if (c3 or c6) else "3col"
                elif len(row) >= 3:
                    col_layout = "3col"
                break
        if col_layout is None:
            col_layout = "3col"

        for row in table:
            if not row:
                continue
            cell0 = str(row[0] or "").strip()
            m = _CRIT_RE.search(cell0)
            if not m:
                continue

            crit_id = m.group(1)
            if crit_id in seen:
                continue
            seen.add(crit_id)

            title = _clean_criterion_title(m.group(2).strip())
            level = m.group(3).upper()

            if col_layout == "9col":
                raw_status = str(row[3] or "").strip() if len(row) > 3 else ""
                remarks_raw = str(row[6] or "").strip() if len(row) > 6 else ""
            else:
                raw_status = str(row[1] or "").strip() if len(row) > 1 else ""
                remarks_raw = str(row[2] or "").strip() if len(row) > 2 else ""

            # Guard against continuation rows being mistaken for remarks
            # If remarks_raw starts with a criterion-like pattern, it's a parsing artifact
            if remarks_raw and _CRIT_RE.match(remarks_raw.strip()):
                remarks_raw = ""

            criteria.append(VPATCriterion(
                criterion_id=crit_id,
                title=title,
                level=level,
                raw_status=raw_status,
                normalized_status=normalize_status(raw_status),
                remarks=_clean_remarks(remarks_raw),
            ))

    return criteria


# ── Text fallback criterion parsing ───────────────────────────────────────────

def _parse_from_text(text: str, seen_ids: set) -> list:
    criteria = []
    lines = text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        m = _CRIT_RE.search(line)
        if m:
            crit_id = m.group(1)
            if crit_id in seen_ids:
                i += 1
                continue
            seen_ids.add(crit_id)

            title = _clean_criterion_title(m.group(2).strip())
            level = m.group(3).upper()

            raw_status, remarks_lines = "", []
            j = i + 1
            found_status = False
            in_remarks = False

            while j < min(i + 25, len(lines)):
                candidate = lines[j].strip()
                if not found_status:
                    sm = _STATUS_RE.match(candidate)
                    if sm:
                        raw_status = sm.group(0)
                        found_status = True
                        j += 1
                        continue
                if found_status:
                    if re.match(r"^Remark", candidate, re.IGNORECASE):
                        in_remarks = True
                        j += 1
                        continue
                    if in_remarks:
                        if _CRIT_RE.search(candidate):
                            break
                        remarks_lines.append(candidate)
                    elif _CRIT_RE.search(candidate):
                        break
                j += 1

            remarks = " ".join(r for r in remarks_lines if r).strip()
            criteria.append(VPATCriterion(
                criterion_id=crit_id,
                title=title,
                level=level,
                raw_status=raw_status,
                normalized_status=normalize_status(raw_status),
                remarks=_clean_remarks(remarks),
            ))
            i = j
        else:
            i += 1
    return criteria


# ── Section 508 parsing ───────────────────────────────────────────────────────

_FPC_TITLES = {
    "302.1": "Without Vision",
    "302.2": "With Limited Vision",
    "302.3": "Without Perception of Color",
    "302.4": "Without Hearing",
    "302.5": "With Limited Hearing",
    "302.6": "Without Speech",
    "302.7": "With Limited Manipulation",
    "302.8": "With Limited Reach and Strength",
    "302.9": "With Limited Language, Cognitive, and Learning Abilities",
}

def _parse_508_fpc(tables: list, text: str) -> list:
    criteria = []
    seen = set()
    fpc_re = re.compile(r"^(302\.\d)\s*(.*)", re.IGNORECASE)

    for table in tables:
        for row in table:
            if not row:
                continue
            cell0 = str(row[0] or "").strip()
            m = fpc_re.match(cell0)
            if not m:
                continue
            cid = m.group(1)
            if cid in seen:
                continue
            seen.add(cid)
            title = _FPC_TITLES.get(cid, m.group(2).strip())

            if len(row) >= 9:
                raw_status, remarks = str(row[3] or "").strip(), str(row[6] or "").strip()
            elif len(row) >= 3:
                raw_status, remarks = str(row[1] or "").strip(), str(row[2] or "").strip()
            elif len(row) >= 2:
                raw_status, remarks = str(row[1] or "").strip(), ""
            else:
                raw_status, remarks = "", ""

            # Skip heading rows
            if "heading cell" in raw_status.lower():
                continue

            criteria.append(VPATCriterion(
                criterion_id=cid, title=title, level="",
                raw_status=raw_status,
                normalized_status=normalize_status(raw_status),
                remarks=_clean_remarks(remarks),
                section="508_fpc",
            ))

    # Text fallback — build all 9 criteria if none found from tables
    if not criteria:
        for cid, ctitle in _FPC_TITLES.items():
            if cid in seen:
                continue
            seen.add(cid)
            # Search for status near criterion id
            pat = re.compile(
                re.escape(cid) + r'.{0,200}?' +
                r'(Supports\s+with\s+Exceptions|Partially\s+Supports?|Does\s+Not\s+Supports?|'
                r'Not\s+Applicable|Not\s+Evaluated|Supports?)',
                re.IGNORECASE | re.DOTALL
            )
            m2 = pat.search(text)
            raw_status = m2.group(1) if m2 else ""
            # Grab remarks after status
            remarks_text = ""
            if m2:
                chunk = text[m2.end():m2.end() + 500]
                stop = re.search(r'302\.\d|Chapter\s+[456]', chunk)
                remarks_text = chunk[:stop.start() if stop else 300]
                remarks_text = re.sub(r'\s+', ' ', remarks_text).strip()
            criteria.append(VPATCriterion(
                criterion_id=cid, title=ctitle, level="",
                raw_status=raw_status,
                normalized_status=normalize_status(raw_status),
                remarks=_clean_remarks(remarks_text),
                section="508_fpc",
            ))

    return criteria


def _parse_508_ch6(tables: list) -> list:
    criteria = []
    seen = set()
    ch6_re = re.compile(r"^(60[23]\.\d)\s*(.*)", re.IGNORECASE)

    for table in tables:
        for row in table:
            if not row:
                continue
            cell0 = str(row[0] or "").strip()
            m = ch6_re.match(cell0)
            if not m:
                continue
            cid = m.group(1)
            title = m.group(2).strip()

            # v9 FIX H: seen.add(cid) was called before the orphan check, so the
            # orphan "602.3 (Support Docs)" row blocked the real "602.3 Electronic
            # Support Documentation" row that comes later in the document.
            # Fix: move seen.add(cid) to after all guards pass.
            # Guard: reject orphan WCAG cross-reference rows BEFORE adding to seen
            # so the real 602.3 row (which comes later in the document) is not blocked
            # e.g. '602.3 (Support Docs)' appears in WCAG table cells, not as a real 602.x entry
            # Real 602.x rows have descriptive titles; orphan rows have '(Support Docs)' etc.
            if re.search(r'\(Support Docs\)|\(Authoring Tool\)|\(Web\)|\(Software\)',
                         m.group(2), re.IGNORECASE):
                continue

            # Passed all guards — now check/update seen set
            if cid in seen:
                continue
            seen.add(cid)

            if len(row) >= 9:
                raw_status, remarks = str(row[3] or "").strip(), str(row[6] or "").strip()
            elif len(row) >= 3:
                raw_status, remarks = str(row[1] or "").strip(), str(row[2] or "").strip()
            else:
                raw_status = str(row[1] or "").strip() if len(row) > 1 else ""
                remarks = ""

            # Skip heading rows
            if "heading cell" in raw_status.lower():
                continue
            normed = normalize_status(raw_status)

            criteria.append(VPATCriterion(
                criterion_id=cid, title=title, level="",
                raw_status=raw_status,
                normalized_status=normed,
                remarks=_clean_remarks(remarks),
                section="508_ch6",
            ))
    return criteria


# ── Main parse entry ──────────────────────────────────────────────────────────

def parse_vpat(filepath: str) -> VPATData:
    data = VPATData()
    ext = Path(filepath).suffix.lower()

    if ext == ".pdf":
        text, tables = _extract_pdf(filepath)
    elif ext in (".docx", ".doc"):
        text, tables = _extract_docx(filepath)
    elif ext == ".txt":
        text, tables = _extract_txt(filepath)
    else:
        data.parse_warnings.append(f"Unsupported file type: {ext}")
        return data

    if not text.strip():
        data.parse_warnings.append("No text could be extracted.")
        return data

    data.raw_text = text

    meta = _extract_meta(text)
    data.product_name         = meta.get("product_name", "")
    data.product_version      = meta.get("product_version", "")
    data.vendor_name          = meta.get("vendor_name", "")
    data.vendor_report_date_raw = meta.get("vendor_report_date_raw", "")
    data.vpat_edition         = meta.get("vpat_edition", "")
    data.vendor_contact       = meta.get("vendor_contact", "")
    data.product_description  = meta.get("product_description", "")
    data.product_type         = meta.get("product_type", "Software")
    data.evaluation_methods   = meta.get("evaluation_methods", "")

    parsed_dt = _parse_date(data.vendor_report_date_raw)
    if parsed_dt:
        data.vendor_report_date = parsed_dt
        delta = (date.today() - parsed_dt).days
        if delta > 365:
            data.is_outdated = True
            months = delta // 30
            data.outdated_note = (
                f"The vendor report date ({data.vendor_report_date_raw}) is approximately "
                f"{months} months old. VPATs older than 12 months may not reflect the current "
                f"accessibility state of the product. SFBRN recommends requesting an updated "
                f"VPAT before final procurement decisions."
            )

    data.standards_reviewed = _detect_standards(text) or []

    table_criteria = _parse_from_tables(tables)
    seen_ids = {c.criterion_id for c in table_criteria}
    text_criteria = _parse_from_text(text, seen_ids)
    data.criteria = table_criteria + text_criteria

    data.criteria += _parse_508_fpc(tables, text)
    data.criteria += _parse_508_ch6(tables)

    logger.info(
        f"Parsed: product='{data.product_name}' vendor='{data.vendor_name}' "
        f"criteria={len(data.criteria)} outdated={data.is_outdated}"
    )
    return data


# ── Analysis helpers ──────────────────────────────────────────────────────────

# Scoring and impact now live in the vpat_reviewer package — policy-driven and
# fully unit-tested (Phase 1 strangler). Defaults reproduce v10 behavior exactly.
# get_aa_barriers keeps its legacy name as an alias of the generic get_barriers.
from vpat_reviewer.domain.impact import calculate_impact
from vpat_reviewer.domain.scoring import compliance_score, get_barriers

get_aa_barriers = get_barriers
